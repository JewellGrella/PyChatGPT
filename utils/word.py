import docx
from PIL import Image
from docx.shared import Cm,  Pt
from docx.oxml.ns import qn
import logging
from utils import ini, file


# 获取docx文件的边距和页面宽度
def get_page_width_and_margins(docx):
    sections = docx.sections
    try:
        if len(sections) > 0:
            section = sections[0]
            page_width = section.page_width.cm
            left_margin = section.left_margin.cm
            right_margin = section.right_margin.cm
            return page_width, left_margin, right_margin
        else:
            return None
    except Exception as e:
        logging.error("获取页面宽度和边距失败，错误信息为：{}".format(e))
        return None


# 在段落中自动插入图片
def add_picture_auto(document, paragraph, picture_path):
    if get_page_width_and_margins(document) is None:
        logging.error("获取页面宽度和边距为None！")
        print("程序退出！")
        exit(1)
    else:
        img = Image.open(picture_path)
        width, height = img.size
        ratio = width / height
        page_width, left_margin, right_margin=get_page_width_and_margins(document)
        new_width = page_width - left_margin - right_margin
        new_height = new_width / ratio
        new_paragraph = paragraph.insert_paragraph_before()
        run = new_paragraph.add_run()
        run.add_picture(picture_path, width=Cm(new_width), height=Cm(new_height))


def save(src_dir, chatbot, dir, file_name):
    # 处理请求的参数
    role = ini.get_config_string("chatgpt", "role")
    if role is None:
        logging.error("获取role失败！")
        print("程序退出！")
        exit(1)
    topic = ini.get_config_string("chatgpt", "topic")
    if topic is None:
        logging.error("获取topic失败！")
        print("程序退出！")
        exit(1)
    # 读取参数文件
    if file.is_dir_exists("static\\params\\{}.txt".format(file_name)):
        with open("static\\params\\{}.txt".format(file_name), 'r', encoding='utf-8') as f:
            params = f.read()
    else:
        logging.error("{}参数文件不存在！".format(file_name))
        print("程序退出！")
        exit(1)
    new_topic = str.replace(str.replace(topic, "?", file_name), "*", params)
    # 发送请求并获得响应数据，这里需要捕获异常
    try:
        content = chatbot.ask(role + "," + new_topic)
    except Exception as e:
        logging.error("发送请求失败，错误信息为：{}".format(e))
        content = None
    if content is None:
        logging.error("获取响应数据失败！")
        print("程序退出！")
        exit(1)
    else:
        # 打开文档
        doc = docx.Document()
        paragraphs = content.split('\n\n')
        # 遍历段落列表，并将每个段落添加到文档中
        for para in paragraphs:
            p = doc.add_paragraph()
            run = p.add_run(para)
            run.font.size = Pt(14)
            run.font.name = u'仿宋'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')

        if file.is_dir_exists(src_dir + "\\" + file_name):
            # 获取图片名称
            pic_names = file.read_file(src_dir + "\\" + file_name, "jpg")
            if pic_names is None:
                logging.error("获取图片名称失败！")
                print("程序退出！")
                exit(1)
            file.is_file_exists_or_create(dir + "\\" + file_name + ".docx")
            paragraphs2 = doc.paragraphs
            # 遍历所有段落
            for p in paragraphs2:
                # 判断段落是否需要插入图片
                if p.text.startswith("外观方面，") or p.text.startswith("外观方面："):
                    add_picture_auto(doc, p, src_dir + "\\"+file_name+"\\" + pic_names[0] + ".jpg")
                    add_picture_auto(doc, p, src_dir+ "\\"+file_name+ "\\" + pic_names[1] + ".jpg")
                if p.text.startswith("内饰方面，") or p.text.startswith("内饰方面："):
                    add_picture_auto(doc, p, src_dir + "\\"+file_name+ "\\" + pic_names[2] + ".jpg")
                if p.text.startswith("动力方面，") or p.text.startswith("动力方面："):
                    add_picture_auto(doc, p, src_dir + "\\"+file_name+ "\\" + pic_names[3] + ".jpg")
                if p.text.startswith("官方指导价，") or p.text.startswith("官方指导价："):
                    add_picture_auto(doc, p, src_dir + "\\"+file_name+"\\" + pic_names[4] + ".jpg")
        # 保存文档
        doc.save(dir + "\\" + file_name + ".docx")

if __name__ == '__main__':
    page_width, left_margin, right_margin = get_page_width_and_margins(
        docx.Document('G:\\Projects\\PyCharm\\py-Chat-GPT\\example.docx'))
    print(page_width, left_margin, right_margin)
